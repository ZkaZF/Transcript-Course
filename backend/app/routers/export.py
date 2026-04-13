"""
Router: Export (Download PDF / Word)
=====================================
Endpoint untuk export transkrip + rangkuman ke PDF atau Word.
"""

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session as DBSession
import io
import os
from datetime import datetime

from ..database import get_db
from ..models import Session as SessionModel, Subject

router = APIRouter()


@router.get("/sessions/{session_id}/export/pdf")
def export_pdf(session_id: int, db: DBSession = Depends(get_db)):
    """Export transkrip + rangkuman sesi ke PDF."""
    session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Sesi tidak ditemukan")
    if not session.transcript and not session.summary:
        raise HTTPException(status_code=400, detail="Belum ada transkrip atau rangkuman untuk di-export")

    subject = db.query(Subject).filter(Subject.id == session.subject_id).first()
    subject_name = subject.name if subject else "Unknown"
    session_title = session.title or f"Pertemuan {session.session_number}"

    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Use built-in Helvetica (supports basic Latin characters)
        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, subject_name, new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, session_title, new_x="LMARGIN", new_y="NEXT")

        # Info
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(100, 100, 100)
        date_str = session.created_at.strftime("%d %B %Y") if session.created_at else "-"
        info_text = f"Pertemuan ke-{session.session_number} | {date_str}"
        if session.duration_minutes > 0:
            info_text += f" | Durasi: {int(session.duration_minutes)} menit"
        pdf.cell(0, 6, info_text, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)

        # Transkrip
        if session.transcript:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "TRANSKRIP", new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(200, 30, 30)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
            pdf.set_font("Helvetica", "", 10)
            # Handle encoding — replace chars not supported by latin-1
            clean_text = session.transcript.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(5)

        # Rangkuman
        if session.summary:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "RANGKUMAN AI", new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(30, 100, 200)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
            pdf.set_font("Helvetica", "", 10)
            clean_summary = session.summary.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, clean_summary)

        # Materi
        if session.material_text:
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "MATERI KULIAH", new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(30, 180, 30)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
            pdf.set_font("Helvetica", "", 10)
            clean_material = session.material_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, clean_material)

        # Generate output
        pdf_output = pdf.output()
        buffer = io.BytesIO(pdf_output)
        buffer.seek(0)

        filename = f"{subject_name} - {session_title}.pdf".replace("/", "-").replace("\\", "-")
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 belum terinstall. Jalankan: pip install fpdf2")


@router.get("/sessions/{session_id}/export/docx")
def export_docx(session_id: int, db: DBSession = Depends(get_db)):
    """Export transkrip + rangkuman sesi ke Word (DOCX)."""
    session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Sesi tidak ditemukan")
    if not session.transcript and not session.summary:
        raise HTTPException(status_code=400, detail="Belum ada transkrip atau rangkuman untuk di-export")

    subject = db.query(Subject).filter(Subject.id == session.subject_id).first()
    subject_name = subject.name if subject else "Unknown"
    session_title = session.title or f"Pertemuan {session.session_number}"

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(subject_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Subtitle
        subtitle = doc.add_heading(session_title, level=1)

        # Info
        date_str = session.created_at.strftime("%d %B %Y") if session.created_at else "-"
        info = f"Pertemuan ke-{session.session_number} | {date_str}"
        if session.duration_minutes > 0:
            info += f" | Durasi: {int(session.duration_minutes)} menit"
        p = doc.add_paragraph(info)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Transkrip
        if session.transcript:
            doc.add_heading("Transkrip", level=2)
            doc.add_paragraph(session.transcript)

        # Rangkuman
        if session.summary:
            doc.add_heading("Rangkuman AI", level=2)
            doc.add_paragraph(session.summary)

        # Materi
        if session.material_text:
            doc.add_heading("Materi Kuliah", level=2)
            doc.add_paragraph(session.material_text)

        # Generate output
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{subject_name} - {session_title}.docx".replace("/", "-").replace("\\", "-")
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx belum terinstall. Jalankan: pip install python-docx")
